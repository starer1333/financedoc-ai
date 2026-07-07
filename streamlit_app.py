from __future__ import annotations

import io
import json
import os
import re
import urllib.error
import urllib.request
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path

import pandas as pd
import plotly.express as px
import streamlit as st
from dotenv import load_dotenv
from PyPDF2 import PdfReader


APP_TITLE = "FinanceDoc AI"
APP_SUBTITLE = "财务文档智能分析与风险审阅系统"
SAMPLE_DATA_PATH = "sample_data/sample_financials.csv"

DEFAULT_SAMPLE_ROWS = [
    {
        "year": 2021,
        "revenue": 82000000,
        "cogs": 49200000,
        "gross_profit": 32800000,
        "operating_expense": 16200000,
        "sales_expense": 5200000,
        "admin_expense": 6100000,
        "r_and_d_expense": 4900000,
        "net_profit": 9600000,
        "operating_cash_flow": 8800000,
        "total_assets": 126000000,
        "total_liabilities": 65000000,
        "equity": 61000000,
        "inventory": 14200000,
        "receivables": 17600000,
    },
    {
        "year": 2022,
        "revenue": 101000000,
        "cogs": 62600000,
        "gross_profit": 38400000,
        "operating_expense": 21300000,
        "sales_expense": 7200000,
        "admin_expense": 7600000,
        "r_and_d_expense": 6500000,
        "net_profit": 11200000,
        "operating_cash_flow": 9400000,
        "total_assets": 151000000,
        "total_liabilities": 80500000,
        "equity": 70500000,
        "inventory": 19800000,
        "receivables": 24600000,
    },
    {
        "year": 2023,
        "revenue": 95000000,
        "cogs": 61750000,
        "gross_profit": 33250000,
        "operating_expense": 26600000,
        "sales_expense": 9000000,
        "admin_expense": 8900000,
        "r_and_d_expense": 8700000,
        "net_profit": 5900000,
        "operating_cash_flow": 3100000,
        "total_assets": 168000000,
        "total_liabilities": 113000000,
        "equity": 55000000,
        "inventory": 27400000,
        "receivables": 32200000,
    },
    {
        "year": 2024,
        "revenue": 108000000,
        "cogs": 68040000,
        "gross_profit": 39960000,
        "operating_expense": 28400000,
        "sales_expense": 9400000,
        "admin_expense": 9300000,
        "r_and_d_expense": 9700000,
        "net_profit": 7800000,
        "operating_cash_flow": 6900000,
        "total_assets": 182000000,
        "total_liabilities": 121000000,
        "equity": 61000000,
        "inventory": 28800000,
        "receivables": 35000000,
    },
]

DEMO_SCENARIOS = {
    "制造业｜利润修复与回款压力": {
        "description": "收入恢复但应收账款和费用投入偏高，适合展示经营质量和现金流审阅。",
        "text": "公司主营高端制造设备，2024 年收入有所恢复，但客户账期拉长，应收账款增加。管理层披露部分大客户验收延迟，存在现金流压力。审计关注收入确认、存货跌价准备、客户集中和关联交易披露。",
        "rows": DEFAULT_SAMPLE_ROWS,
    },
    "SaaS｜高增长与亏损收窄": {
        "description": "收入高速增长、研发投入高、经营现金流改善，适合展示成长型企业分析。",
        "text": "公司提供企业级 SaaS 服务，订阅收入增长较快，续费率稳定。报告期内研发费用保持高位，销售费用率下降，经营现金流由负转正。风险关注客户集中、收入递延确认、资本化研发支出和数据合规。",
        "rows": [
            {"year": 2021, "revenue": 36000000, "cogs": 9000000, "gross_profit": 27000000, "operating_expense": 42000000, "sales_expense": 16000000, "admin_expense": 8200000, "r_and_d_expense": 17800000, "net_profit": -15200000, "operating_cash_flow": -11800000, "total_assets": 78000000, "total_liabilities": 36000000, "equity": 42000000, "inventory": 0, "receivables": 7200000},
            {"year": 2022, "revenue": 59000000, "cogs": 13600000, "gross_profit": 45400000, "operating_expense": 56000000, "sales_expense": 21800000, "admin_expense": 9600000, "r_and_d_expense": 24600000, "net_profit": -8900000, "operating_cash_flow": -4200000, "total_assets": 104000000, "total_liabilities": 47000000, "equity": 57000000, "inventory": 0, "receivables": 11800000},
            {"year": 2023, "revenue": 91000000, "cogs": 20000000, "gross_profit": 71000000, "operating_expense": 76000000, "sales_expense": 28600000, "admin_expense": 12400000, "r_and_d_expense": 35000000, "net_profit": -2600000, "operating_cash_flow": 1800000, "total_assets": 138000000, "total_liabilities": 62000000, "equity": 76000000, "inventory": 0, "receivables": 16400000},
            {"year": 2024, "revenue": 128000000, "cogs": 28160000, "gross_profit": 99840000, "operating_expense": 93000000, "sales_expense": 32600000, "admin_expense": 14500000, "r_and_d_expense": 45900000, "net_profit": 4200000, "operating_cash_flow": 9600000, "total_assets": 176000000, "total_liabilities": 76500000, "equity": 99500000, "inventory": 0, "receivables": 22600000},
        ],
    },
    "零售连锁｜库存周转与费用压力": {
        "description": "收入小幅下滑、库存和费用承压，适合展示风险识别和周转分析。",
        "text": "公司经营区域零售连锁门店，受客流下降和促销影响，收入下降、毛利率下降。报告期末存货增加，部分门店关闭并计提减值。风险关注存货跌价、租赁负债、费用控制和门店关闭损失。",
        "rows": [
            {"year": 2021, "revenue": 148000000, "cogs": 96200000, "gross_profit": 51800000, "operating_expense": 38600000, "sales_expense": 25200000, "admin_expense": 9000000, "r_and_d_expense": 0, "net_profit": 8600000, "operating_cash_flow": 10200000, "total_assets": 116000000, "total_liabilities": 59000000, "equity": 57000000, "inventory": 28800000, "receivables": 8600000},
            {"year": 2022, "revenue": 156000000, "cogs": 103000000, "gross_profit": 53000000, "operating_expense": 41400000, "sales_expense": 27600000, "admin_expense": 9600000, "r_and_d_expense": 0, "net_profit": 7100000, "operating_cash_flow": 6600000, "total_assets": 124000000, "total_liabilities": 68400000, "equity": 55600000, "inventory": 34200000, "receivables": 9200000},
            {"year": 2023, "revenue": 142000000, "cogs": 98700000, "gross_profit": 43300000, "operating_expense": 42800000, "sales_expense": 28900000, "admin_expense": 10100000, "r_and_d_expense": 0, "net_profit": 1200000, "operating_cash_flow": -1800000, "total_assets": 121000000, "total_liabilities": 75200000, "equity": 45800000, "inventory": 39800000, "receivables": 9800000},
            {"year": 2024, "revenue": 135000000, "cogs": 95850000, "gross_profit": 39150000, "operating_expense": 43600000, "sales_expense": 29600000, "admin_expense": 10400000, "r_and_d_expense": 0, "net_profit": -3900000, "operating_cash_flow": -5200000, "total_assets": 113000000, "total_liabilities": 81500000, "equity": 31500000, "inventory": 42600000, "receivables": 10400000},
        ],
    },
    "地产服务｜高负债与偿债压力": {
        "description": "收入下降、负债率高、现金流紧张，适合展示审计风险和持续经营分析。",
        "text": "公司为地产项目提供管理和配套服务，受行业下行影响，收入下降，回款周期延长。存在债务逾期、担保、诉讼和持续经营重大不确定性。审计重点关注资产减值、关联方资金占用、偿债压力和现金流预测。",
        "rows": [
            {"year": 2021, "revenue": 210000000, "cogs": 147000000, "gross_profit": 63000000, "operating_expense": 39800000, "sales_expense": 8600000, "admin_expense": 23800000, "r_and_d_expense": 0, "net_profit": 15400000, "operating_cash_flow": 9600000, "total_assets": 420000000, "total_liabilities": 278000000, "equity": 142000000, "inventory": 76000000, "receivables": 62000000},
            {"year": 2022, "revenue": 188000000, "cogs": 139000000, "gross_profit": 49000000, "operating_expense": 41200000, "sales_expense": 8200000, "admin_expense": 25600000, "r_and_d_expense": 0, "net_profit": 5200000, "operating_cash_flow": 1800000, "total_assets": 438000000, "total_liabilities": 318000000, "equity": 120000000, "inventory": 93000000, "receivables": 82000000},
            {"year": 2023, "revenue": 151000000, "cogs": 119000000, "gross_profit": 32000000, "operating_expense": 38600000, "sales_expense": 7600000, "admin_expense": 24800000, "r_and_d_expense": 0, "net_profit": -11600000, "operating_cash_flow": -18400000, "total_assets": 402000000, "total_liabilities": 334000000, "equity": 68000000, "inventory": 101000000, "receivables": 96000000},
            {"year": 2024, "revenue": 132000000, "cogs": 109000000, "gross_profit": 23000000, "operating_expense": 36200000, "sales_expense": 7100000, "admin_expense": 23800000, "r_and_d_expense": 0, "net_profit": -21800000, "operating_cash_flow": -29600000, "total_assets": 368000000, "total_liabilities": 326000000, "equity": 42000000, "inventory": 106000000, "receivables": 104000000},
        ],
    },
    "拟 IPO｜收入增长与关联交易审阅": {
        "description": "高增长但关联交易和客户集中明显，适合展示 IPO 尽调视角。",
        "text": "公司处于拟 IPO 阶段，收入增长较快，但前五大客户占比较高。报告期内存在关联交易、关联方采购和担保事项。审阅重点包括收入真实性、关联交易定价公允性、客户集中、供应商集中和内控有效性。",
        "rows": [
            {"year": 2021, "revenue": 68000000, "cogs": 40800000, "gross_profit": 27200000, "operating_expense": 13800000, "sales_expense": 4200000, "admin_expense": 5600000, "r_and_d_expense": 4000000, "net_profit": 8400000, "operating_cash_flow": 6200000, "total_assets": 96000000, "total_liabilities": 42000000, "equity": 54000000, "inventory": 12600000, "receivables": 15600000},
            {"year": 2022, "revenue": 96000000, "cogs": 56640000, "gross_profit": 39360000, "operating_expense": 18600000, "sales_expense": 6200000, "admin_expense": 7200000, "r_and_d_expense": 5200000, "net_profit": 14200000, "operating_cash_flow": 10600000, "total_assets": 128000000, "total_liabilities": 56200000, "equity": 71800000, "inventory": 17800000, "receivables": 22800000},
            {"year": 2023, "revenue": 138000000, "cogs": 80040000, "gross_profit": 57960000, "operating_expense": 26800000, "sales_expense": 9100000, "admin_expense": 9200000, "r_and_d_expense": 8500000, "net_profit": 22600000, "operating_cash_flow": 17100000, "total_assets": 176000000, "total_liabilities": 78600000, "equity": 97400000, "inventory": 25200000, "receivables": 34600000},
            {"year": 2024, "revenue": 182000000, "cogs": 105560000, "gross_profit": 76440000, "operating_expense": 35600000, "sales_expense": 12400000, "admin_expense": 11600000, "r_and_d_expense": 11600000, "net_profit": 31400000, "operating_cash_flow": 23800000, "total_assets": 238000000, "total_liabilities": 105000000, "equity": 133000000, "inventory": 35600000, "receivables": 51200000},
        ],
    },
}

LLM_PROVIDERS = {
    "Demo Mode": {
        "api_key_env": "",
        "base_url_env": "",
        "model_env": "",
        "default_base_url": "",
        "default_model": "",
    },
    "OpenAI": {
        "api_key_env": "OPENAI_API_KEY",
        "base_url_env": "OPENAI_BASE_URL",
        "model_env": "OPENAI_MODEL",
        "default_base_url": "https://api.openai.com/v1/chat/completions",
        "default_model": "gpt-4o-mini",
    },
    "DeepSeek": {
        "api_key_env": "DEEPSEEK_API_KEY",
        "base_url_env": "DEEPSEEK_BASE_URL",
        "model_env": "DEEPSEEK_MODEL",
        "default_base_url": "https://api.deepseek.com/chat/completions",
        "default_model": "deepseek-chat",
    },
}


FINANCIAL_ALIASES = {
    "year": ["year", "年份", "年度", "period", "报告期"],
    "revenue": ["revenue", "营业收入", "收入", "sales", "operating revenue"],
    "cogs": ["cogs", "营业成本", "成本", "cost of revenue", "cost"],
    "gross_profit": ["gross_profit", "毛利", "毛利润", "gross profit"],
    "operating_expense": ["operating_expense", "期间费用", "运营费用", "经营费用", "opex"],
    "sales_expense": ["sales_expense", "销售费用", "selling expense", "sales expense"],
    "admin_expense": ["admin_expense", "管理费用", "administrative expense"],
    "r_and_d_expense": ["r_and_d_expense", "研发费用", "research", "r&d"],
    "net_profit": ["net_profit", "净利润", "profit", "net income", "净收益"],
    "operating_cash_flow": ["operating_cash_flow", "经营现金流", "经营活动现金流", "ocf"],
    "total_assets": ["total_assets", "总资产", "资产总计", "assets"],
    "total_liabilities": ["total_liabilities", "总负债", "负债合计", "liabilities"],
    "equity": ["equity", "所有者权益", "股东权益", "shareholder equity"],
    "inventory": ["inventory", "存货"],
    "receivables": ["receivables", "应收账款", "应收款"],
}

RISK_KEYWORDS = {
    "审计意见风险": ["保留意见", "否定意见", "无法表示意见", "强调事项", "持续经营重大不确定性"],
    "合规风险": ["处罚", "诉讼", "违规", "监管问询", "立案调查"],
    "关联交易风险": ["关联交易", "关联方", "资金占用", "担保"],
    "流动性风险": ["债务逾期", "流动性紧张", "现金流压力", "偿债压力"],
    "经营风险": ["收入下降", "毛利率下降", "客户集中", "供应商集中", "减值"],
}


@dataclass
class RiskItem:
    category: str
    level: str
    finding: str
    suggestion: str


def configure_page() -> None:
    load_dotenv()
    st.set_page_config(page_title=f"{APP_TITLE} | {APP_SUBTITLE}", page_icon="FD", layout="wide")
    st.markdown(
        """
        <style>
        .main .block-container {
            padding-top: 1.25rem;
            padding-bottom: 3rem;
            max-width: 1240px;
        }
        h1, h2, h3 {
            letter-spacing: 0;
        }
        h1 {
            font-size: clamp(2rem, 4vw, 3.4rem) !important;
            line-height: 1.08 !important;
        }
        h2 {
            font-size: 1.45rem !important;
            margin-top: 1.4rem !important;
        }
        h3 {
            font-size: 1.08rem !important;
        }
        .metric-row [data-testid="stMetric"] {
            background: #ffffff;
            border: 1px solid #e5e7eb;
            border-radius: 8px;
            padding: 14px 16px;
        }
        .small-note {
            color: #5b6472;
            font-size: 0.92rem;
            line-height: 1.55;
        }
        .section-band {
            border-left: 4px solid #1f6f8b;
            padding: 0.75rem 1rem;
            background: #f7fafc;
            margin: 0.5rem 0 1rem 0;
        }
        .hero {
            padding: 0.35rem 0 0.75rem 0;
            border-bottom: 1px solid #e5e7eb;
            margin-bottom: 1rem;
        }
        .hero-kicker {
            color: #1f6f8b;
            font-size: 0.88rem;
            font-weight: 700;
            letter-spacing: 0.04em;
            text-transform: uppercase;
        }
        .hero-copy {
            color: #4b5563;
            font-size: 1rem;
            line-height: 1.7;
            max-width: 900px;
        }
        .section-anchor {
            scroll-margin-top: 1.5rem;
        }
        .status-ok {
            color: #047857;
            font-weight: 700;
        }
        .status-warn {
            color: #b45309;
            font-weight: 700;
        }
        </style>
        """,
        unsafe_allow_html=True,
    )


def read_sample_data() -> pd.DataFrame:
    sample_path = Path(SAMPLE_DATA_PATH)
    if sample_path.exists():
        return pd.read_csv(sample_path)
    return pd.DataFrame(DEFAULT_SAMPLE_ROWS)


def scenario_dataframe(name: str) -> pd.DataFrame:
    scenario = DEMO_SCENARIOS.get(name, DEMO_SCENARIOS["制造业｜利润修复与回款压力"])
    return pd.DataFrame(scenario["rows"])


def scenario_text(name: str) -> str:
    scenario = DEMO_SCENARIOS.get(name, DEMO_SCENARIOS["制造业｜利润修复与回款压力"])
    return str(scenario["text"])


def apply_demo_scenario(name: str) -> None:
    df = scenario_dataframe(name)
    st.session_state.document_text = scenario_text(name)
    st.session_state.tables = {name: df}
    st.session_state.active_df = df
    st.session_state.source_name = f"内置样例：{name}"
    st.session_state.custom_uploaded = False
    st.session_state.demo_scenario = name


def provider_is_configured(provider: str) -> bool:
    provider_config = LLM_PROVIDERS.get(provider, LLM_PROVIDERS["Demo Mode"])
    api_key_env = provider_config["api_key_env"]
    if not api_key_env:
        return True
    return bool(get_secret(api_key_env))


def provider_status_label(provider: str) -> str:
    if provider == "Demo Mode":
        return "Demo Mode：已启用，本地 RAG + 规则分析"
    provider_config = LLM_PROVIDERS[provider]
    api_key_env = provider_config["api_key_env"]
    if provider_is_configured(provider):
        return f"{provider}：已配置 {api_key_env}，可调用真实模型"
    return f"{provider}：未配置 {api_key_env}，会自动回退 Demo Mode"


def normalize_name(name: object) -> str:
    return re.sub(r"\s+", " ", str(name).strip().lower())


def find_column(df: pd.DataFrame, key: str) -> str | None:
    normalized_columns = {normalize_name(column): column for column in df.columns}
    for alias in FINANCIAL_ALIASES.get(key, []):
        alias_normalized = normalize_name(alias)
        if alias_normalized in normalized_columns:
            return normalized_columns[alias_normalized]
    for column in df.columns:
        column_normalized = normalize_name(column)
        if any(normalize_name(alias) in column_normalized for alias in FINANCIAL_ALIASES.get(key, [])):
            return column
    return None


def numeric_series(df: pd.DataFrame, key: str) -> pd.Series | None:
    column = find_column(df, key)
    if column is None:
        return None
    series = pd.to_numeric(df[column], errors="coerce")
    return series if series.notna().any() else None


def active_year_series(df: pd.DataFrame) -> pd.Series:
    year_column = find_column(df, "year")
    if year_column:
        return df[year_column].astype(str)
    return pd.Series([f"Period {idx + 1}" for idx in range(len(df))])


def fmt_money(value: float | int | None) -> str:
    if value is None or pd.isna(value):
        return "N/A"
    abs_value = abs(float(value))
    if abs_value >= 100_000_000:
        return f"{value / 100_000_000:.2f} 亿"
    if abs_value >= 10_000:
        return f"{value / 10_000:.2f} 万"
    return f"{value:,.0f}"


def fmt_pct(value: float | None) -> str:
    if value is None or pd.isna(value):
        return "N/A"
    return f"{value * 100:.1f}%"


def pct_change(series: pd.Series | None) -> float | None:
    if series is None or len(series.dropna()) < 2:
        return None
    clean = series.dropna()
    previous = clean.iloc[-2]
    latest = clean.iloc[-1]
    if previous == 0:
        return None
    return (latest - previous) / abs(previous)


def latest_value(series: pd.Series | None) -> float | None:
    if series is None or series.dropna().empty:
        return None
    return float(series.dropna().iloc[-1])


def extract_pdf_text(file: io.BytesIO) -> str:
    reader = PdfReader(file)
    pages = []
    for index, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text() or ""
        if page_text.strip():
            pages.append(f"[第 {index} 页]\n{page_text.strip()}")
    return "\n\n".join(pages)


def parse_uploaded_file(uploaded_file) -> tuple[str, dict[str, pd.DataFrame]]:
    suffix = uploaded_file.name.split(".")[-1].lower()
    raw = uploaded_file.getvalue()
    if suffix == "pdf":
        text = extract_pdf_text(io.BytesIO(raw))
        return text, {}
    if suffix in {"xlsx", "xls"}:
        sheets = pd.read_excel(io.BytesIO(raw), sheet_name=None)
        text_parts = [f"Sheet: {name}\n{frame.head(30).to_markdown(index=False)}" for name, frame in sheets.items()]
        return "\n\n".join(text_parts), sheets
    if suffix == "csv":
        df = pd.read_csv(io.BytesIO(raw))
        return df.head(80).to_markdown(index=False), {"CSV": df}
    raise ValueError("暂不支持该文件类型，请上传 PDF、Excel 或 CSV。")


def parse_uploaded_files(uploaded_files) -> tuple[str, dict[str, pd.DataFrame], list[str]]:
    text_parts: list[str] = []
    tables: dict[str, pd.DataFrame] = {}
    names: list[str] = []
    for uploaded_file in uploaded_files:
        text, file_tables = parse_uploaded_file(uploaded_file)
        names.append(uploaded_file.name)
        text_parts.append(f"## 文件：{uploaded_file.name}\n{text}")
        for table_name, frame in file_tables.items():
            tables[f"{uploaded_file.name} / {table_name}"] = frame
    return "\n\n".join(text_parts), tables, names


def pick_primary_table(tables: dict[str, pd.DataFrame]) -> pd.DataFrame:
    if not tables:
        return read_sample_data()
    scored: list[tuple[int, str, pd.DataFrame]] = []
    for name, frame in tables.items():
        financial_columns = sum(1 for key in FINANCIAL_ALIASES if find_column(frame, key))
        numeric_columns = len(frame.select_dtypes(include="number").columns)
        scored.append((financial_columns * 3 + numeric_columns, name, frame))
    scored.sort(key=lambda item: item[0], reverse=True)
    return scored[0][2].copy()


def build_financial_summary(df: pd.DataFrame) -> dict[str, str]:
    revenue = numeric_series(df, "revenue")
    cogs = numeric_series(df, "cogs")
    gross_profit = numeric_series(df, "gross_profit")
    operating_expense = numeric_series(df, "operating_expense")
    net_profit = numeric_series(df, "net_profit")
    operating_cash_flow = numeric_series(df, "operating_cash_flow")
    total_assets = numeric_series(df, "total_assets")
    total_liabilities = numeric_series(df, "total_liabilities")

    latest_revenue = latest_value(revenue)
    latest_net_profit = latest_value(net_profit)
    latest_ocf = latest_value(operating_cash_flow)
    latest_assets = latest_value(total_assets)
    latest_liabilities = latest_value(total_liabilities)
    latest_gross_profit = latest_value(gross_profit)

    if latest_gross_profit is None and latest_revenue is not None and latest_value(cogs) is not None:
        latest_gross_profit = latest_revenue - latest_value(cogs)

    gross_margin = latest_gross_profit / latest_revenue if latest_revenue else None
    net_margin = latest_net_profit / latest_revenue if latest_revenue else None
    liability_ratio = latest_liabilities / latest_assets if latest_assets else None
    cash_conversion = latest_ocf / latest_net_profit if latest_net_profit else None
    expense_ratio = latest_value(operating_expense) / latest_revenue if latest_revenue and latest_value(operating_expense) else None

    return {
        "latest_revenue": fmt_money(latest_revenue),
        "revenue_growth": fmt_pct(pct_change(revenue)),
        "latest_net_profit": fmt_money(latest_net_profit),
        "net_profit_growth": fmt_pct(pct_change(net_profit)),
        "gross_margin": fmt_pct(gross_margin),
        "net_margin": fmt_pct(net_margin),
        "operating_cash_flow": fmt_money(latest_ocf),
        "cash_conversion": fmt_pct(cash_conversion),
        "liability_ratio": fmt_pct(liability_ratio),
        "expense_ratio": fmt_pct(expense_ratio),
    }


def scan_text_risks(text: str) -> list[RiskItem]:
    findings: list[RiskItem] = []
    for category, keywords in RISK_KEYWORDS.items():
        matched = [word for word in keywords if word in text]
        if matched:
            findings.append(
                RiskItem(
                    category=category,
                    level="中",
                    finding=f"文档中出现关键词：{', '.join(matched[:5])}",
                    suggestion="建议回到原文定位上下文，核查金额、期间、管理层解释及后续整改情况。",
                )
            )
    return findings


def detect_financial_risks(df: pd.DataFrame, text: str) -> list[RiskItem]:
    risks: list[RiskItem] = []
    revenue_growth = pct_change(numeric_series(df, "revenue"))
    net_profit_growth = pct_change(numeric_series(df, "net_profit"))
    ocf = latest_value(numeric_series(df, "operating_cash_flow"))
    net_profit = latest_value(numeric_series(df, "net_profit"))
    revenue = latest_value(numeric_series(df, "revenue"))
    operating_expense = latest_value(numeric_series(df, "operating_expense"))
    assets = latest_value(numeric_series(df, "total_assets"))
    liabilities = latest_value(numeric_series(df, "total_liabilities"))
    receivables_growth = pct_change(numeric_series(df, "receivables"))
    inventory_growth = pct_change(numeric_series(df, "inventory"))

    if revenue_growth is not None and revenue_growth < -0.08:
        risks.append(
            RiskItem("收入趋势", "高", f"营业收入同比变化为 {fmt_pct(revenue_growth)}，存在明显下滑。", "拆分产品、客户和地区，判断是周期性波动还是竞争力下降。")
        )
    if net_profit_growth is not None and net_profit_growth < -0.15:
        risks.append(
            RiskItem("盈利能力", "高", f"净利润同比变化为 {fmt_pct(net_profit_growth)}，利润承压。", "进一步分析毛利率、费用率、减值和非经常性损益。")
        )
    if net_profit and ocf is not None and ocf < net_profit * 0.6:
        risks.append(
            RiskItem("现金流质量", "中", f"经营现金流为 {fmt_money(ocf)}，显著低于净利润 {fmt_money(net_profit)}。", "重点查看应收账款、存货、预收款和现金回款节奏。")
        )
    if assets and liabilities and liabilities / assets > 0.7:
        risks.append(
            RiskItem("资产负债结构", "高", f"资产负债率约为 {fmt_pct(liabilities / assets)}。", "关注短期借款、到期债务、利息保障倍数和再融资安排。")
        )
    if revenue and operating_expense and operating_expense / revenue > 0.28:
        risks.append(
            RiskItem("费用控制", "中", f"期间费用率约为 {fmt_pct(operating_expense / revenue)}。", "拆分销售、管理、研发费用，确认费用投入是否能转化为收入增长。")
        )
    if receivables_growth is not None and revenue_growth is not None and receivables_growth - revenue_growth > 0.15:
        risks.append(
            RiskItem("回款质量", "中", f"应收账款增速 {fmt_pct(receivables_growth)} 高于收入增速 {fmt_pct(revenue_growth)}。", "关注客户信用政策变化、账龄结构和坏账准备充分性。")
        )
    if inventory_growth is not None and revenue_growth is not None and inventory_growth - revenue_growth > 0.15:
        risks.append(
            RiskItem("存货周转", "中", f"存货增速 {fmt_pct(inventory_growth)} 高于收入增速 {fmt_pct(revenue_growth)}。", "关注跌价准备、备货策略和产品滞销风险。")
        )

    risks.extend(scan_text_risks(text))
    if not risks:
        risks.append(
            RiskItem("综合判断", "低", "基于当前 demo 规则，未发现明显高风险信号。", "仍建议结合原始报表附注、审计意见和行业数据交叉验证。")
        )
    return risks


def split_text(text: str, chunk_size: int = 520) -> list[str]:
    paragraphs = [item.strip() for item in re.split(r"\n{2,}|。|；", text) if item.strip()]
    chunks: list[str] = []
    current = ""
    for paragraph in paragraphs:
        if len(current) + len(paragraph) > chunk_size and current:
            chunks.append(current)
            current = paragraph
        else:
            current = f"{current}。{paragraph}" if current else paragraph
    if current:
        chunks.append(current)
    return chunks


def table_context(df: pd.DataFrame) -> str:
    summary = build_financial_summary(df)
    return "\n".join(
        [
            f"最新收入：{summary['latest_revenue']}",
            f"收入增速：{summary['revenue_growth']}",
            f"最新净利润：{summary['latest_net_profit']}",
            f"净利润增速：{summary['net_profit_growth']}",
            f"毛利率：{summary['gross_margin']}",
            f"净利率：{summary['net_margin']}",
            f"经营现金流：{summary['operating_cash_flow']}",
            f"现金利润转化率：{summary['cash_conversion']}",
            f"资产负债率：{summary['liability_ratio']}",
            f"费用率：{summary['expense_ratio']}",
        ]
    )


def retrieve_context(question: str, text: str, df: pd.DataFrame, top_k: int = 4) -> list[str]:
    chunks = split_text(text)
    chunks.append("财务指标摘要：\n" + table_context(df))
    terms = [term for term in re.split(r"\W+|的|和|是|吗|如何|为什么|请|分析", question) if len(term) >= 2]
    scored: list[tuple[int, int, str]] = []
    for index, chunk in enumerate(chunks):
        lowered = chunk.lower()
        score = sum(lowered.count(term.lower()) for term in terms)
        score += sum(1 for term in terms if term.lower() in lowered)
        if score or "财务指标摘要" in chunk:
            scored.append((score, -index, chunk))
    scored.sort(key=lambda item: (item[0], item[1]), reverse=True)
    return [chunk for _, _, chunk in scored[:top_k]]


def get_secret(name: str) -> str | None:
    try:
        value = st.secrets.get(name)
        if value:
            return str(value)
    except Exception:  # noqa: BLE001
        pass
    return os.getenv(name)


def call_llm(provider: str, prompt: str) -> tuple[str, bool]:
    provider_config = LLM_PROVIDERS.get(provider, LLM_PROVIDERS["Demo Mode"])
    api_key_env = provider_config["api_key_env"]
    if not api_key_env:
        return "", False

    api_key = get_secret(api_key_env)
    if not api_key:
        return f"未配置 {api_key_env}，当前自动使用 demo mode 回答。", False

    base_url = get_secret(provider_config["base_url_env"]) or provider_config["default_base_url"]
    model = get_secret(provider_config["model_env"]) or provider_config["default_model"]
    payload = {
        "model": model,
        "messages": [
            {"role": "system", "content": "你是严谨的中文财务分析助手。回答必须基于用户提供的上下文，不能编造数据。"},
            {"role": "user", "content": prompt},
        ],
        "temperature": 0.2,
    }
    request = urllib.request.Request(
        base_url,
        data=json.dumps(payload).encode("utf-8"),
        headers={"Content-Type": "application/json", "Authorization": f"Bearer {api_key}"},
        method="POST",
    )
    try:
        with urllib.request.urlopen(request, timeout=25) as response:
            result = json.loads(response.read().decode("utf-8"))
        return result["choices"][0]["message"]["content"], True
    except (urllib.error.URLError, urllib.error.HTTPError, KeyError, IndexError, json.JSONDecodeError) as exc:
        return f"{provider} 调用失败，已回退 demo mode。错误信息：{exc}", False


def answer_question(question: str, df: pd.DataFrame, text: str, risks: list[RiskItem]) -> str:
    question = question.strip()
    if not question:
        return "请输入一个具体问题，例如：收入为什么下降？现金流质量如何？有哪些审计风险？"

    summary = build_financial_summary(df)
    lower_question = question.lower()
    if any(word in question for word in ["风险", "问题", "异常", "审计"]):
        risk_lines = [f"- {item.category}（{item.level}）：{item.finding}" for item in risks[:5]]
        return "基于当前规则识别，主要风险包括：\n\n" + "\n".join(risk_lines)
    if any(word in question for word in ["收入", "营收", "revenue"]):
        return f"当前表格中最新收入为 {summary['latest_revenue']}，收入增速为 {summary['revenue_growth']}。建议进一步拆分客户、产品和地区，判断增长质量。"
    if any(word in question for word in ["利润", "盈利", "profit"]):
        return f"最新净利润为 {summary['latest_net_profit']}，净利润增速为 {summary['net_profit_growth']}，净利率为 {summary['net_margin']}。若利润与收入趋势背离，应重点查看毛利率和费用率。"
    if any(word in question for word in ["现金流", "回款", "cash"]):
        return f"经营现金流为 {summary['operating_cash_flow']}，现金利润转化率为 {summary['cash_conversion']}。如果该指标低于 100%，需要关注应收账款、存货和付款周期。"
    if "api" in lower_question or "llm" in lower_question or "deepseek" in lower_question or "openai" in lower_question:
        return "第一版采用 demo mode：规则分析和文本检索优先，保证可部署和可解释。后续可以在同一接口层接入 OpenAI 或 DeepSeek，把检索到的上下文交给模型生成更自然的答案。"

    query_terms = [term for term in re.split(r"\W+|的|和|是|吗|如何|为什么", question) if len(term) >= 2]
    chunks = split_text(text)
    scored: list[tuple[int, str]] = []
    for chunk in chunks:
        score = sum(chunk.lower().count(term.lower()) for term in query_terms)
        if score:
            scored.append((score, chunk))
    scored.sort(key=lambda item: item[0], reverse=True)
    if scored:
        evidence = "\n\n".join(f"- {chunk[:280]}" for _, chunk in scored[:3])
        return f"我在已解析文本中找到了这些相关片段，可作为初步回答依据：\n\n{evidence}\n\n建议结合原文页码和表格数据进一步复核。"
    return "当前文档中没有检索到强相关片段。你可以换一个更贴近财务字段或风险关键词的问题。"


def answer_question_with_rag(question: str, df: pd.DataFrame, text: str, risks: list[RiskItem], provider: str) -> str:
    contexts = retrieve_context(question, text, df)
    risk_lines = "\n".join(f"- {item.category}（{item.level}）：{item.finding}" for item in risks[:6])
    prompt = f"""请基于以下资料回答财务分析问题。

问题：{question}

检索上下文：
{chr(10).join(f'[{idx + 1}] {chunk}' for idx, chunk in enumerate(contexts))}

规则识别风险：
{risk_lines}

回答要求：
1. 先给结论。
2. 引用关键数据或文本证据。
3. 给出下一步审阅建议。
4. 如果资料不足，明确说明不足。
"""
    llm_answer, used_llm = call_llm(provider, prompt)
    if used_llm:
        return llm_answer

    fallback = answer_question(question, df, text, risks)
    evidence = "\n\n".join(f"- {chunk[:320]}" for chunk in contexts)
    prefix = f"{llm_answer}\n\n" if llm_answer else ""
    return f"{prefix}{fallback}\n\n**RAG 检索证据：**\n\n{evidence}"


def create_brief(df: pd.DataFrame, text: str, risks: list[RiskItem]) -> str:
    summary = build_financial_summary(df)
    risk_lines = "\n".join([f"- {item.category}：{item.finding}" for item in risks[:6]])
    keyword_hits = []
    for category, keywords in RISK_KEYWORDS.items():
        matched = [word for word in keywords if word in text]
        if matched:
            keyword_hits.append(f"{category}：{', '.join(matched[:4])}")
    keyword_section = "\n".join(f"- {item}" for item in keyword_hits) if keyword_hits else "- 未在文本中识别到高频异常关键词。"

    return f"""# FinanceDoc AI 分析底稿

生成时间：{datetime.now().strftime("%Y-%m-%d %H:%M")}

## 1. 商业与财务概览

- 最新营业收入：{summary["latest_revenue"]}
- 收入增速：{summary["revenue_growth"]}
- 最新净利润：{summary["latest_net_profit"]}
- 净利润增速：{summary["net_profit_growth"]}
- 毛利率：{summary["gross_margin"]}
- 净利率：{summary["net_margin"]}

## 2. 现金流与资产负债

- 经营现金流：{summary["operating_cash_flow"]}
- 现金利润转化率：{summary["cash_conversion"]}
- 资产负债率：{summary["liability_ratio"]}
- 费用率：{summary["expense_ratio"]}

## 3. 主要风险点

{risk_lines}

## 4. 文本关键词信号

{keyword_section}

## 5. 面试或投研追问

- 收入增长或下滑主要来自价格、销量、客户结构还是行业周期？
- 净利润变化是否由主营业务驱动，还是由减值、投资收益或非经常性项目驱动？
- 经营现金流与净利润是否匹配，应收账款和存货是否出现异常累积？
- 审计意见、关联交易、担保、诉讼和监管问询是否存在重大影响？
- 如果接入 LLM API，哪些段落需要进入 RAG 检索上下文以提高回答可信度？
"""


def run_multi_agent_review(df: pd.DataFrame, text: str, risks: list[RiskItem], provider: str = "Demo Mode") -> dict[str, str]:
    summary = build_financial_summary(df)
    risk_lines = "\n".join(f"- {item.category}（{item.level}）：{item.finding}" for item in risks[:6])
    context = "\n".join(retrieve_context("财务表现 风险 审计 现金流 经营策略", text, df, top_k=5))
    prompt = f"""请用四个角色分析财务文档：

财务指标：
{table_context(df)}

主要风险：
{risk_lines}

文档上下文：
{context}

请分别输出：
Accounting Agent、Risk Review Agent、Strategy Analyst Agent、Report Writer Agent。
"""
    llm_answer, used_llm = call_llm(provider, prompt)
    if used_llm:
        return {"LLM Multi-Agent Review": llm_answer}

    return {
        "Accounting Agent": (
            f"最新收入为 {summary['latest_revenue']}，收入增速 {summary['revenue_growth']}；"
            f"最新净利润为 {summary['latest_net_profit']}，净利率 {summary['net_margin']}。"
            "建议继续核对收入确认政策、成本结转、费用归集和非经常性损益。"
        ),
        "Risk Review Agent": (
            f"当前识别到 {len(risks)} 条风险信号。重点关注："
            + "；".join(f"{item.category}: {item.finding}" for item in risks[:4])
        ),
        "Strategy Analyst Agent": (
            "从经营视角看，需要判断收入恢复是否来自核心业务，同时关注费用投入、客户结构、库存和应收账款对现金流的影响。"
        ),
        "Report Writer Agent": (
            "建议报告结构采用：业务概览、关键财务表现、现金流和资产负债、核心风险、管理层追问、后续审阅清单。"
        ),
    }


def create_audit_workpaper(df: pd.DataFrame, text: str, risks: list[RiskItem], agent_outputs: dict[str, str]) -> str:
    brief = create_brief(df, text, risks)
    agents = "\n\n".join(f"## {name}\n\n{content}" for name, content in agent_outputs.items())
    checklist = "\n".join(
        [
            "- 核对收入确认政策与主要合同条款。",
            "- 抽查大额应收账款回款记录及期后回款。",
            "- 复核存货跌价准备和库龄结构。",
            "- 检查关联交易定价公允性和披露完整性。",
            "- 对高风险事项补充管理层访谈问题。",
        ]
    )
    return f"""{brief}

## 6. Multi-Agent 审阅意见

{agents}

## 7. 审计底稿检查清单

{checklist}
"""


def create_docx_bytes(markdown_text: str) -> bytes:
    from docx import Document

    document = Document()
    for line in markdown_text.splitlines():
        if line.startswith("# "):
            document.add_heading(line[2:].strip(), level=1)
        elif line.startswith("## "):
            document.add_heading(line[3:].strip(), level=2)
        elif line.startswith("- "):
            document.add_paragraph(line[2:].strip(), style="List Bullet")
        elif line.strip():
            document.add_paragraph(line.strip())
    buffer = io.BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def create_pdf_bytes(markdown_text: str) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.pdfbase import pdfmetrics
    from reportlab.pdfbase.cidfonts import UnicodeCIDFont
    from reportlab.pdfgen import canvas

    buffer = io.BytesIO()
    pdfmetrics.registerFont(UnicodeCIDFont("STSong-Light"))
    page = canvas.Canvas(buffer, pagesize=A4)
    width, height = A4
    x = 42
    y = height - 48
    page.setFont("STSong-Light", 10)
    for raw_line in markdown_text.splitlines():
        line = raw_line.replace("#", "").strip()
        if not line:
            y -= 8
            continue
        wrapped = [line[index : index + 42] for index in range(0, len(line), 42)]
        for part in wrapped:
            if y < 48:
                page.showPage()
                page.setFont("STSong-Light", 10)
                y = height - 48
            page.drawString(x, y, part)
            y -= 16
    page.save()
    return buffer.getvalue()


def ensure_session_state() -> None:
    default_scenario = "制造业｜利润修复与回款压力"
    if "demo_scenario" not in st.session_state:
        st.session_state.demo_scenario = default_scenario
    if "custom_uploaded" not in st.session_state:
        st.session_state.custom_uploaded = False
    if "document_text" not in st.session_state:
        st.session_state.document_text = scenario_text(default_scenario)
    if "tables" not in st.session_state:
        st.session_state.tables = {default_scenario: scenario_dataframe(default_scenario)}
    if "active_df" not in st.session_state:
        st.session_state.active_df = scenario_dataframe(default_scenario)
    if "source_name" not in st.session_state:
        st.session_state.source_name = f"内置样例：{default_scenario}"


def render_header() -> None:
    st.title(f"{APP_TITLE}｜{APP_SUBTITLE}")
    st.caption("面向财务分析、审计、投研和商业分析场景的文档解析与风险审阅 demo。")


def page_home() -> None:
    render_header()
    st.markdown(
        """
        <div class="section-band">
        FinanceDoc AI 将“财务文档上传、文本和表格解析、关键指标提取、规则化风险识别、分析底稿生成”串成一个可演示的工作流。
        第一版不依赖本地 Ollama 或付费 API，因此适合部署到 Streamlit Community Cloud。
        </div>
        """,
        unsafe_allow_html=True,
    )
    col1, col2, col3 = st.columns(3)
    col1.metric("输入文件", "PDF / Excel / CSV")
    col2.metric("分析模式", "RAG + Multi-Agent")
    col3.metric("部署方式", "Streamlit Cloud")

    st.subheader("核心场景")
    st.write(
        "上传年报、审计报告、招股书或财务分析表后，系统会提取可读文本和表格数据，自动生成财务摘要、风险审阅、问答结果和分析师 brief。"
    )

    st.subheader("技术栈")
    st.write("Streamlit、Pandas、PyPDF2、openpyxl、Plotly、python-docx、ReportLab。Q&A 可选择 demo mode、OpenAI 或 DeepSeek。")


def page_upload_parse() -> None:
    render_header()
    uploaded_files = st.file_uploader("上传 PDF、Excel 或 CSV 文件", type=["pdf", "xlsx", "xls", "csv"], accept_multiple_files=True)
    if uploaded_files:
        try:
            text, tables, names = parse_uploaded_files(uploaded_files)
            st.session_state.document_text = text or "文件已解析，但没有提取到可读文本。"
            st.session_state.tables = tables or {"Demo": read_sample_data()}
            st.session_state.active_df = pick_primary_table(st.session_state.tables)
            st.session_state.source_name = "；".join(names)
            st.success(f"已解析 {len(names)} 个文件：{'、'.join(names)}")
        except Exception as exc:  # noqa: BLE001
            st.error(f"解析失败：{exc}")

    st.info(f"当前数据源：{st.session_state.source_name}")
    tab1, tab2 = st.tabs(["文本预览", "表格预览"])
    with tab1:
        st.text_area("已提取文本", value=st.session_state.document_text[:6000], height=360)
    with tab2:
        for name, frame in st.session_state.tables.items():
            st.markdown(f"**{name}**")
            st.dataframe(frame.head(100), width="stretch")


def page_financial_summary() -> None:
    render_header()
    df = st.session_state.active_df
    summary = build_financial_summary(df)
    st.markdown('<div class="metric-row">', unsafe_allow_html=True)
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("最新收入", summary["latest_revenue"], summary["revenue_growth"])
    col2.metric("最新净利润", summary["latest_net_profit"], summary["net_profit_growth"])
    col3.metric("毛利率", summary["gross_margin"])
    col4.metric("资产负债率", summary["liability_ratio"])
    st.markdown("</div>", unsafe_allow_html=True)

    years = active_year_series(df)
    revenue = numeric_series(df, "revenue")
    net_profit = numeric_series(df, "net_profit")
    operating_cash_flow = numeric_series(df, "operating_cash_flow")
    operating_expense = numeric_series(df, "operating_expense")

    chart_col1, chart_col2 = st.columns(2)
    with chart_col1:
        if revenue is not None or net_profit is not None:
            chart_df = pd.DataFrame({"期间": years})
            if revenue is not None:
                chart_df["营业收入"] = revenue
            if net_profit is not None:
                chart_df["净利润"] = net_profit
            melted = chart_df.melt(id_vars="期间", var_name="指标", value_name="金额")
            fig = px.line(melted, x="期间", y="金额", color="指标", markers=True, title="收入与利润趋势")
            st.plotly_chart(fig, width="stretch")
    with chart_col2:
        values = {
            "期间费用": latest_value(operating_expense),
            "经营现金流": latest_value(operating_cash_flow),
            "净利润": latest_value(net_profit),
        }
        bar_df = pd.DataFrame({"指标": list(values.keys()), "金额": list(values.values())}).dropna()
        if not bar_df.empty:
            fig = px.bar(bar_df, x="指标", y="金额", title="利润、现金流与费用对比", color="指标")
            st.plotly_chart(fig, width="stretch")

    st.subheader("数据表")
    st.dataframe(df, width="stretch")


def page_risk_review() -> None:
    render_header()
    risks = detect_financial_risks(st.session_state.active_df, st.session_state.document_text)
    risk_df = pd.DataFrame([item.__dict__ for item in risks])
    st.dataframe(risk_df, width="stretch", hide_index=True)

    level_order = ["低", "中", "高"]
    count_df = risk_df.groupby("level", as_index=False).size()
    count_df["level"] = pd.Categorical(count_df["level"], categories=level_order, ordered=True)
    count_df = count_df.sort_values("level")
    fig = px.bar(count_df, x="level", y="size", title="风险等级分布", labels={"level": "风险等级", "size": "数量"}, color="level")
    st.plotly_chart(fig, width="stretch")

    st.subheader("审阅建议")
    for item in risks:
        with st.expander(f"{item.level}风险｜{item.category}", expanded=item.level == "高"):
            st.write(item.finding)
            st.write(item.suggestion)


def page_qa() -> None:
    render_header()
    st.write("支持 RAG 检索问答；配置 API key 后可调用 OpenAI 或 DeepSeek，未配置时自动回退 demo mode。")
    provider = st.selectbox("回答模式", list(LLM_PROVIDERS.keys()))
    question = st.text_input("输入你的问题", placeholder="例如：这家公司现金流质量怎么样？有哪些审计风险？")
    if st.button("生成回答", type="primary"):
        risks = detect_financial_risks(st.session_state.active_df, st.session_state.document_text)
        st.markdown(answer_question_with_rag(question, st.session_state.active_df, st.session_state.document_text, risks, provider))

    with st.expander("API 配置说明"):
        st.write("OpenAI：配置 `OPENAI_API_KEY`，可选 `OPENAI_MODEL` 和 `OPENAI_BASE_URL`。")
        st.write("DeepSeek：配置 `DEEPSEEK_API_KEY`，可选 `DEEPSEEK_MODEL` 和 `DEEPSEEK_BASE_URL`。")


def page_rag_search() -> None:
    render_header()
    query = st.text_input("检索问题或关键词", value="现金流 风险 审计意见")
    top_k = st.slider("返回片段数量", min_value=2, max_value=8, value=4)
    contexts = retrieve_context(query, st.session_state.document_text, st.session_state.active_df, top_k=top_k)
    st.subheader("检索结果")
    for index, chunk in enumerate(contexts, start=1):
        with st.expander(f"证据片段 {index}", expanded=index == 1):
            st.write(chunk)


def page_multi_agent() -> None:
    render_header()
    provider = st.selectbox("分析模式", list(LLM_PROVIDERS.keys()), key="agent_provider")
    risks = detect_financial_risks(st.session_state.active_df, st.session_state.document_text)
    agent_outputs = run_multi_agent_review(st.session_state.active_df, st.session_state.document_text, risks, provider)
    for name, content in agent_outputs.items():
        with st.expander(name, expanded=True):
            st.write(content)


def page_analyst_brief() -> None:
    render_header()
    risks = detect_financial_risks(st.session_state.active_df, st.session_state.document_text)
    brief = create_brief(st.session_state.active_df, st.session_state.document_text, risks)
    st.markdown(brief)
    st.download_button("下载 Markdown 分析底稿", brief, file_name="financedoc_ai_brief.md", mime="text/markdown")


def page_export() -> None:
    render_header()
    risks = detect_financial_risks(st.session_state.active_df, st.session_state.document_text)
    agent_outputs = run_multi_agent_review(st.session_state.active_df, st.session_state.document_text, risks)
    brief = create_audit_workpaper(st.session_state.active_df, st.session_state.document_text, risks, agent_outputs)
    risk_csv = pd.DataFrame([item.__dict__ for item in risks]).to_csv(index=False).encode("utf-8-sig")
    data_csv = st.session_state.active_df.to_csv(index=False).encode("utf-8-sig")
    docx_bytes = create_docx_bytes(brief)
    pdf_bytes = create_pdf_bytes(brief)

    col1, col2, col3 = st.columns(3)
    col1.download_button("下载分析底稿", brief, file_name="finance_analysis_brief.md", mime="text/markdown")
    col2.download_button("下载风险清单 CSV", risk_csv, file_name="risk_review.csv", mime="text/csv")
    col3.download_button("下载解析表格 CSV", data_csv, file_name="parsed_financial_table.csv", mime="text/csv")
    col4, col5 = st.columns(2)
    col4.download_button("下载 Word 审计底稿", docx_bytes, file_name="audit_workpaper.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    col5.download_button("下载 PDF 审计底稿", pdf_bytes, file_name="audit_workpaper.pdf", mime="application/pdf")

    st.subheader("导出预览")
    st.markdown(brief)


def render_hero(provider: str) -> None:
    status_class = "status-ok" if provider_is_configured(provider) else "status-warn"
    st.markdown(
        f"""
        <div class="hero section-anchor" id="overview">
            <div class="hero-kicker">Public Streamlit App · Finance + AI Portfolio Project</div>
            <h1>{APP_TITLE}｜{APP_SUBTITLE}</h1>
            <p class="hero-copy">
            一个面向财务分析、审计、投研和商业分析场景的网页应用。打开页面后可直接向下滚动，
            完成“上传/选择样例数据 → 财务摘要 → 风险审阅 → RAG 问答 → Multi-Agent 分析 → 导出底稿”的完整工作流。
            </p>
            <p class="{status_class}">当前 AI 模式：{provider_status_label(provider)}</p>
        </div>
        """,
        unsafe_allow_html=True,
    )
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("支持文件", "PDF / Excel / CSV")
    col2.metric("内置样例", f"{len(DEMO_SCENARIOS)} 套")
    col3.metric("分析框架", "RAG + Agents")
    col4.metric("导出格式", "MD / CSV / Word / PDF")


def render_data_section() -> None:
    st.markdown('<div class="section-anchor" id="data"></div>', unsafe_allow_html=True)
    st.header("1. 数据输入与样例场景")
    st.write("可上传真实文件，也可以直接切换内置行业样例。未上传时，系统使用当前样例数据进行完整分析。")

    uploaded_files = st.file_uploader(
        "上传 PDF、Excel 或 CSV 文件",
        type=["pdf", "xlsx", "xls", "csv"],
        accept_multiple_files=True,
        key="one_page_uploader",
    )
    if uploaded_files:
        try:
            text, tables, names = parse_uploaded_files(uploaded_files)
            st.session_state.document_text = text or "文件已解析，但没有提取到可读文本。"
            st.session_state.tables = tables or {st.session_state.demo_scenario: scenario_dataframe(st.session_state.demo_scenario)}
            st.session_state.active_df = pick_primary_table(st.session_state.tables)
            st.session_state.source_name = "；".join(names)
            st.session_state.custom_uploaded = True
            st.success(f"已解析 {len(names)} 个文件：{'、'.join(names)}")
        except Exception as exc:  # noqa: BLE001
            st.error(f"解析失败：{exc}")

    if st.button("恢复当前内置样例数据"):
        apply_demo_scenario(st.session_state.demo_scenario)
        st.rerun()

    st.info(f"当前数据源：{st.session_state.source_name}")
    scenario = DEMO_SCENARIOS.get(st.session_state.demo_scenario)
    if scenario and not st.session_state.custom_uploaded:
        st.caption(scenario["description"])

    tab1, tab2 = st.tabs(["表格数据", "文本上下文"])
    with tab1:
        for name, frame in st.session_state.tables.items():
            st.markdown(f"**{name}**")
            st.dataframe(frame.head(100), width="stretch")
    with tab2:
        st.text_area("已解析文本或样例业务描述", value=st.session_state.document_text[:6000], height=220)


def render_summary_section() -> None:
    st.markdown('<div class="section-anchor" id="summary"></div>', unsafe_allow_html=True)
    st.header("2. 财务摘要与趋势")
    df = st.session_state.active_df
    summary = build_financial_summary(df)
    col1, col2, col3, col4 = st.columns(4)
    col1.metric("最新收入", summary["latest_revenue"], summary["revenue_growth"])
    col2.metric("最新净利润", summary["latest_net_profit"], summary["net_profit_growth"])
    col3.metric("毛利率", summary["gross_margin"])
    col4.metric("资产负债率", summary["liability_ratio"])

    years = active_year_series(df)
    revenue = numeric_series(df, "revenue")
    net_profit = numeric_series(df, "net_profit")
    operating_cash_flow = numeric_series(df, "operating_cash_flow")
    operating_expense = numeric_series(df, "operating_expense")

    chart_col1, chart_col2 = st.columns(2)
    with chart_col1:
        chart_df = pd.DataFrame({"期间": years})
        if revenue is not None:
            chart_df["营业收入"] = revenue
        if net_profit is not None:
            chart_df["净利润"] = net_profit
        if len(chart_df.columns) > 1:
            melted = chart_df.melt(id_vars="期间", var_name="指标", value_name="金额")
            fig = px.line(melted, x="期间", y="金额", color="指标", markers=True, title="收入与利润趋势")
            st.plotly_chart(fig, width="stretch")
    with chart_col2:
        values = {
            "期间费用": latest_value(operating_expense),
            "经营现金流": latest_value(operating_cash_flow),
            "净利润": latest_value(net_profit),
        }
        bar_df = pd.DataFrame({"指标": list(values.keys()), "金额": list(values.values())}).dropna()
        if not bar_df.empty:
            fig = px.bar(bar_df, x="指标", y="金额", title="利润、现金流与费用对比", color="指标")
            st.plotly_chart(fig, width="stretch")


def render_risk_section() -> list[RiskItem]:
    st.markdown('<div class="section-anchor" id="risk"></div>', unsafe_allow_html=True)
    st.header("3. 风险审阅")
    risks = detect_financial_risks(st.session_state.active_df, st.session_state.document_text)
    risk_df = pd.DataFrame([item.__dict__ for item in risks])
    col1, col2 = st.columns([1.15, 0.85])
    with col1:
        st.dataframe(risk_df, width="stretch", hide_index=True)
    with col2:
        count_df = risk_df.groupby("level", as_index=False).size()
        level_order = ["低", "中", "高"]
        count_df["level"] = pd.Categorical(count_df["level"], categories=level_order, ordered=True)
        count_df = count_df.sort_values("level")
        fig = px.bar(count_df, x="level", y="size", title="风险等级分布", labels={"level": "风险等级", "size": "数量"}, color="level")
        st.plotly_chart(fig, width="stretch")
    for item in risks[:4]:
        with st.expander(f"{item.level}风险｜{item.category}", expanded=item.level == "高"):
            st.write(item.finding)
            st.write(item.suggestion)
    return risks


def render_qa_section(provider: str, risks: list[RiskItem]) -> None:
    st.markdown('<div class="section-anchor" id="qa"></div>', unsafe_allow_html=True)
    st.header("4. RAG 问答助手")
    st.write("提问时系统会先检索当前文档和表格摘要；如果已配置 API key，会把检索证据交给 OpenAI 或 DeepSeek 生成更专业的回答。")
    question = st.text_input("输入问题", value="这家公司最需要关注的财务风险是什么？")
    if st.button("生成 RAG 回答", type="primary"):
        st.markdown(answer_question_with_rag(question, st.session_state.active_df, st.session_state.document_text, risks, provider))
    with st.expander("查看当前问题的检索证据"):
        contexts = retrieve_context(question, st.session_state.document_text, st.session_state.active_df, top_k=4)
        for index, chunk in enumerate(contexts, start=1):
            st.markdown(f"**证据 {index}**")
            st.write(chunk)


def render_agents_section(provider: str, risks: list[RiskItem]) -> dict[str, str]:
    st.markdown('<div class="section-anchor" id="agents"></div>', unsafe_allow_html=True)
    st.header("5. Multi-Agent 专业审阅")
    st.write("这里用四个专业角色模拟财务分析、风险审阅、战略分析和报告写作。配置 API 后，同一套上下文会交给真实模型生成。")
    agent_outputs = run_multi_agent_review(st.session_state.active_df, st.session_state.document_text, risks, provider)
    cols = st.columns(2)
    for index, (name, content) in enumerate(agent_outputs.items()):
        with cols[index % 2]:
            with st.expander(name, expanded=True):
                st.write(content)
    return agent_outputs


def render_export_section(risks: list[RiskItem], agent_outputs: dict[str, str]) -> None:
    st.markdown('<div class="section-anchor" id="export"></div>', unsafe_allow_html=True)
    st.header("6. 分析底稿导出")
    brief = create_audit_workpaper(st.session_state.active_df, st.session_state.document_text, risks, agent_outputs)
    risk_csv = pd.DataFrame([item.__dict__ for item in risks]).to_csv(index=False).encode("utf-8-sig")
    data_csv = st.session_state.active_df.to_csv(index=False).encode("utf-8-sig")
    docx_bytes = create_docx_bytes(brief)
    pdf_bytes = create_pdf_bytes(brief)

    col1, col2, col3, col4, col5 = st.columns(5)
    col1.download_button("Markdown", brief, file_name="finance_analysis_brief.md", mime="text/markdown")
    col2.download_button("风险 CSV", risk_csv, file_name="risk_review.csv", mime="text/csv")
    col3.download_button("表格 CSV", data_csv, file_name="parsed_financial_table.csv", mime="text/csv")
    col4.download_button("Word", docx_bytes, file_name="audit_workpaper.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
    col5.download_button("PDF", pdf_bytes, file_name="audit_workpaper.pdf", mime="application/pdf")

    with st.expander("预览分析底稿", expanded=False):
        st.markdown(brief)


def render_api_section(provider: str) -> None:
    st.markdown('<div class="section-anchor" id="api"></div>', unsafe_allow_html=True)
    st.header("7. API 配置状态")
    status_rows = []
    for name in LLM_PROVIDERS:
        status_rows.append(
            {
                "模式": name,
                "状态": provider_status_label(name),
                "当前选择": "是" if name == provider else "否",
            }
        )
    st.dataframe(pd.DataFrame(status_rows), width="stretch", hide_index=True)
    st.info(
        "如果未配置 API key，网站仍会用 RAG 检索 + 财务规则回答问题；配置 OpenAI 或 DeepSeek 后，回答会基于同样证据交给模型生成，适合对外展示更自然、更专业的分析。"
    )


def render_one_page(provider: str) -> None:
    render_hero(provider)
    render_data_section()
    render_summary_section()
    risks = render_risk_section()
    render_qa_section(provider, risks)
    agent_outputs = render_agents_section(provider, risks)
    render_export_section(risks, agent_outputs)
    render_api_section(provider)


def render_sidebar_settings() -> str:
    st.sidebar.title(APP_TITLE)
    st.sidebar.caption("公网作品集版本")
    st.sidebar.markdown(
        "[概览](#overview) · [数据](#data) · [摘要](#summary) · [风险](#risk) · [问答](#qa) · [Agent](#agents) · [导出](#export)"
    )
    provider = st.sidebar.selectbox("AI 模式", list(LLM_PROVIDERS.keys()))
    scenario_names = list(DEMO_SCENARIOS.keys())
    selected_scenario = st.sidebar.selectbox(
        "内置样例数据",
        scenario_names,
        index=scenario_names.index(st.session_state.demo_scenario),
    )
    if selected_scenario != st.session_state.demo_scenario and not st.session_state.custom_uploaded:
        apply_demo_scenario(selected_scenario)
        st.rerun()
    if selected_scenario != st.session_state.demo_scenario and st.session_state.custom_uploaded:
        st.session_state.demo_scenario = selected_scenario
    if st.sidebar.button("载入所选样例"):
        apply_demo_scenario(selected_scenario)
        st.rerun()
    st.sidebar.divider()
    for name in LLM_PROVIDERS:
        st.sidebar.caption(provider_status_label(name))
    return provider


def render_sidebar() -> str:
    st.sidebar.title(APP_TITLE)
    st.sidebar.caption(APP_SUBTITLE)
    return st.sidebar.radio(
        "导航",
        [
            "Home",
            "Upload & Parse",
            "Financial Summary",
            "Risk Review",
            "Q&A Assistant",
            "RAG Search",
            "Multi-Agent Review",
            "Analyst Brief",
            "Export",
        ],
    )


def main() -> None:
    configure_page()
    ensure_session_state()
    provider = render_sidebar_settings()
    render_one_page(provider)


if __name__ == "__main__":
    main()
